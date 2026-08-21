"""
scripts/seed_knowledge_base.py
------------------------------
Ingest the product documentation folder (docs/) into the AI knowledge base
tables so the chatbot's RAG pipeline can answer questions grounded in it:

    docs/*.md|*.docx  ──►  ai_knowledge_source / ai_knowledge_document /
                           ai_knowledge_chunk   (namespace: billing_public)

Idempotent: a document is re-ingested only when its content hash changed
(old versions are superseded, never deleted, per ZB-AI-KB-001 versioning).

Usage:
    python -m scripts.seed_knowledge_base [--docs-dir ../docs] [--recreate]
"""

import argparse
import hashlib
import sys
from datetime import datetime, timezone
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from app.database import SessionLocal, initialize_database
from app.modules.chatbot.models import (
    KnowledgeNamespace,
    KnowledgeSource,
    KnowledgeDocument,
    KnowledgeChunk,
    KnowledgeClassification,
    KnowledgeSourceDocType,
    FreshnessStatus,
)

NAMESPACE_CODE = "billing_public"
SOURCE_TITLE = "Zoiko Billing Product Docs"
CHUNK_TARGET = 700
CHUNK_MAX = 900

# Scope guard (ZB-AI-KB-001): billing_public is the USER-FACING chatbot KB.
# Only files explicitly allowlisted here may ever be indexed — everything
# else in docs/ is internal engineering material (FRS / PRD / API /
# architecture / DB-schema / guardrail wireframes) and must stay out of the
# reachable index entirely, not merely be filtered at query time.
#
# The allowlist is deliberately EMPTY: docs/ currently contains no
# user-facing documents. The live chat KB is seeded by backend/seed_knowledge.py
# ("Zoiko Billing Knowledge Base" source). To publish a genuinely public doc,
# add its filename pattern here AND get product/compliance sign-off.
PUBLIC_DOC_PATTERNS: tuple[str, ...] = ()


def _is_public_doc(path: Path) -> bool:
    name = path.name.lower()
    return any(p in name for p in PUBLIC_DOC_PATTERNS)


# ── Text extraction ──────────────────────────────────────────────────────────

def extract_markdown(path: Path) -> tuple[str, list[tuple[str | None, str]]]:
    """Return (title, [(section_heading|None, text)]). Sections split on '#'."""
    text = path.read_text(encoding="utf-8", errors="replace")
    title = None
    sections: list[tuple[str | None, list[str]]] = []
    current_heading = None
    current: list[str] = []
    for line in text.splitlines():
        m = line.startswith("#")
        if m:
            if title is None and line.lstrip().startswith("# "):
                title = line.lstrip()[2:].strip()
            if current:
                sections.append((current_heading, "\n".join(current).strip()))
                current = []
            current_heading = line.lstrip("#").strip()
        else:
            current.append(line)
    if current:
        sections.append((current_heading, "\n".join(current).strip()))
    return (
        title or path.stem.replace("_", " ").replace("-", " ").title(),
        [(h, t) for h, t in sections if t],
    )


def extract_docx(path: Path) -> tuple[str, list[tuple[str | None, str]]]:
    """Return (title, [(section_heading|None, text)]) using python-docx.
    Tables are flattened to pipe-separated rows.

    Word core-properties titles are notoriously generic ('Word Document',
    'CHATBOT') and collide across files, so the FILENAME stem is always the
    canonical document title."""
    import docx  # python-docx

    d = docx.Document(str(path))
    sections: list[tuple[str | None, list[str]]] = []
    current_heading = None
    current: list[str] = []

    def flush():
        if current:
            sections.append((current_heading, "\n".join(current).strip()))
            current.clear()

    for para in d.paragraphs:
        text = (para.text or "").strip()
        style = (para.style.name if para.style is not None else "") or ""
        if style.lower().startswith("heading") or style.lower() == "title":
            flush()
            current_heading = text or None
        elif text:
            current.append(text)
        # empty paragraphs act as soft separators but don't split sections
    for table in d.tables:
        rows = [
            " | ".join((cell.text or "").strip() for cell in row.cells)
            for row in table.rows
        ]
        rows = [r for r in rows if r.strip(" |")]
        if rows:
            flush()
            sections.append((current_heading, "\n".join(rows)))
    return (
        path.stem.replace("_", " ").replace("-", " ").strip(),
        [(h, t) for h, t in sections if t],
    )


EXTRACTORS = {
    ".md": extract_markdown,
    ".markdown": extract_markdown,
    ".docx": extract_docx,
    ".txt": lambda p: (p.stem, [(None, p.read_text(encoding="utf-8", errors="replace"))]),
}


# ── Chunking ─────────────────────────────────────────────────────────────────

def _split_long(text: str) -> list[str]:
    if len(text) <= CHUNK_MAX:
        return [text]
    parts, buf = [], ""
    for sentence in text.replace("! ", ". ").replace("? ", ". ").split(". "):
        piece = sentence if sentence.endswith(".") else sentence + "."
        if len(buf) + len(piece) + 1 > CHUNK_TARGET and buf:
            parts.append(buf.strip())
            buf = piece
        else:
            buf = f"{buf} {piece}".strip()
    if buf.strip():
        parts.append(buf.strip())
    return parts


def pack_chunks(sections: list[tuple[str | None, str]]) -> list[str]:
    """Pack section paragraphs into retrieval-sized chunks (~700 chars),
    keeping each section's heading attached to its content."""
    chunks: list[str] = []
    for heading, body in sections:
        pieces = [p.strip() for p in body.split("\n") if p.strip()]
        buf = f"{heading}\n" if heading else ""
        for piece in pieces:
            for sub in _split_long(piece):
                candidate = f"{buf} {sub}".strip() if buf else sub
                if len(candidate) > CHUNK_MAX and buf:
                    chunks.append(buf.strip())
                    buf = sub
                else:
                    buf = candidate
        if buf.strip():
            chunks.append(buf.strip())
    return [c for c in chunks if c]


# ── Ingestion ────────────────────────────────────────────────────────────────

def _ensure_namespace(db) -> KnowledgeNamespace:
    ns = db.query(KnowledgeNamespace).filter(
        KnowledgeNamespace.namespace_code == NAMESPACE_CODE
    ).first()
    if not ns:
        ns = KnowledgeNamespace(
            namespace_code=NAMESPACE_CODE,
            tenant_id=0,
            allowed_domains='["billing","help","dashboard"]',
            description="Public Zoiko Billing product documentation",
        )
        db.add(ns)
        db.flush()
    return ns


def _ensure_source(db, ns_id: int) -> KnowledgeSource:
    src = db.query(KnowledgeSource).filter(
        KnowledgeSource.namespace_id == ns_id,
        KnowledgeSource.title == SOURCE_TITLE,
    ).first()
    if not src:
        src = KnowledgeSource(
            namespace_id=ns_id,
            source_type=KnowledgeSourceDocType.DOC,
            classification=KnowledgeClassification.INTERNAL,
            owner_team="product",
            title=SOURCE_TITLE,
            source_url=None,
            status="active",
        )
        db.add(src)
        db.flush()
    return src


def ingest_file(db, src_id: int, path: Path, force: bool) -> str:
    extractor = EXTRACTORS.get(path.suffix.lower())
    if extractor is None:
        return "skipped (unsupported type)"
    title, sections = extractor(path)
    chunks = pack_chunks(sections)
    if not chunks:
        return "skipped (no extractable text)"
    # Hash covers title + content so a title fix re-ingests the file.
    content_hash = hashlib.sha256(
        (title + "\n" + "\n".join(chunks)).encode("utf-8")
    ).hexdigest()

    # Document identity is the SOURCE FILE (object_uri), never the title —
    # Word core titles are unreliable and can collide across files.
    existing = db.query(KnowledgeDocument).filter(
        KnowledgeDocument.source_id == src_id,
        KnowledgeDocument.status == "approved",
        KnowledgeDocument.object_uri == str(path),
    ).all()
    if existing:
        if not force and any(d.document_hash == content_hash for d in existing):
            return "unchanged"
        # Retire old versions in ALL cases (force included) — leaving them
        # approved duplicates every chunk in the retrieval index.
        for d in existing:
            d.status = "superseded"
            d.freshness_status = FreshnessStatus.EXPIRED
            d.superseded_at = datetime.now(timezone.utc)

    doc = KnowledgeDocument(
        source_id=src_id,
        document_version=(max((d.document_version or 1) for d in existing) + 1) if existing else 1,
        document_hash=content_hash,
        freshness_status=FreshnessStatus.CURRENT,
        object_uri=str(path),
        title=title,
        status="approved",
        approved_at=datetime.now(timezone.utc),
    )
    db.add(doc)
    db.flush()
    for seq, text in enumerate(chunks, 1):
        db.add(KnowledgeChunk(
            document_id=doc.id,
            chunk_sequence=seq,
            chunk_text=text,
            classification=KnowledgeClassification.INTERNAL,
        ))
    return f"ingested v{doc.document_version} ({len(chunks)} chunks)"


def retire_nonpublic_docs(db, src_id: int) -> int:
    """Supersede any approved docs in this source that are not on the public
    allowlist — e.g. indexed by an older seeder version before the gate.
    This is the INDEX-level exclusion: superseded docs are unreachable by
    the retriever regardless of query wording."""
    retired = 0
    for d in db.query(KnowledgeDocument).filter(
        KnowledgeDocument.source_id == src_id,
        KnowledgeDocument.status == "approved",
    ).all():
        uri = (d.object_uri or "").replace("\\", "/").lower()
        if not _is_public_doc(Path(uri)):
            d.status = "superseded"
            d.freshness_status = FreshnessStatus.EXPIRED
            d.superseded_at = datetime.now(timezone.utc)
            retired += 1
    return retired


def main() -> None:
    parser = argparse.ArgumentParser(description="Seed the AI knowledge base from docs/.")
    here = Path(__file__).resolve().parent.parent  # backend/
    parser.add_argument(
        "--docs-dir", default=str(here.parent / "docs"),
        help="Folder containing .md/.docx product docs",
    )
    parser.add_argument("--recreate", action="store_true",
                        help="Re-ingest even when the content hash is unchanged")
    args = parser.parse_args()

    initialize_database()
    docs_dir = Path(args.docs_dir).resolve()
    files = sorted(
        p for p in docs_dir.rglob("*")
        if p.is_file() and p.suffix.lower() in EXTRACTORS
        and _is_public_doc(p)
    )
    if not files:
        print(
            f"No public-allowlisted documents found in {docs_dir} "
            f"(allowlist: {PUBLIC_DOC_PATTERNS or 'EMPTY — docs/ is internal-only'})"
        )

    db = SessionLocal()
    try:
        ns = _ensure_namespace(db)
        src = _ensure_source(db, ns.id)
        print(f"Namespace '{NAMESPACE_CODE}' (id={ns.id}), source '{SOURCE_TITLE}' (id={src.id})")
        for path in files:
            rel = path.relative_to(docs_dir)
            try:
                status_line = ingest_file(db, src.id, path, args.recreate)
            except Exception as exc:  # noqa: BLE001 — report and continue
                db.rollback()
                status_line = f"FAILED: {exc}"
            print(f"  {rel}: {status_line}")
        retired = retire_nonpublic_docs(db, src.id)
        if retired:
            print(f"  scope guard: retired {retired} non-public doc(s) from the index")
        db.commit()
        print("Done.")
    finally:
        db.close()


if __name__ == "__main__":
    main()
